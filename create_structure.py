#!/usr/bin/env python3
"""
УНИВЕРСАЛЬНЫЙ генератор структуры документов из DOCX файлов
Автоматически определяет тип и извлекает ВСЕ структурные элементы
"""

import re
import logging
from pathlib import Path
from typing import List, Dict, Any
from docx import Document

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class UniversalStructureExtractor:
    """Универсальный извлекатель структуры для ЛЮБЫХ нормативных актов"""

    def __init__(self, docx_path: Path):
        self.docx_path = docx_path
        self.doc = Document(str(docx_path))
        self.paragraphs = [p for p in self.doc.paragraphs if p.text.strip()]
        self.doc_type = self._detect_type()

    def _detect_type(self) -> str:
        """Автоопределение типа документа"""
        name = self.docx_path.stem.lower()

        if 'кодекс' in name or 'жилищный кодекс' in name:
            return 'CODE'
        if 'фз' in name or ('№' in name and 'федеральный закон' in name):
            return 'CODE'
        if 'пп рф' in name or 'постановление' in name:
            return 'GOVERNMENT_DECREE'
        if 'приказ' in name:
            return 'MINISTRY_ORDER'
        if 'письмо' in name:
            return 'LETTER'
        return 'UNKNOWN'

    def extract_structure(self) -> Dict[str, Any]:
        """Универсальное извлечение структуры"""

        if self.doc_type == 'CODE':
            return self._extract_code_structure()
        elif self.doc_type == 'GOVERNMENT_DECREE':
            return self._extract_decree_structure()
        elif self.doc_type == 'MINISTRY_ORDER':
            return self._extract_order_structure()
        else:
            return self._extract_simple_structure()

    def _extract_code_structure(self) -> Dict[str, Any]:
        """Структура кодекса/ФЗ: Разделы -> Главы -> Статьи"""
        structure = {
            'type': 'CODE',
            'sections': [],
            'total_paragraphs': len(self.paragraphs),
            'total_tables': len(self.doc.tables)
        }

        current_section = None
        current_chapter = None

        for i, p in enumerate(self.paragraphs):
            text = p.text.strip()

            # Раздел I, II, III
            if re.match(r'^Раздел\s+[IVXLCDM]+', text, re.IGNORECASE):
                current_section = {
                    'title': text,
                    'index': i,
                    'chapters': []
                }
                structure['sections'].append(current_section)
                current_chapter = None
                continue

            # Глава 1, 2, 3
            if re.match(r'^Глава\s+\d+', text, re.IGNORECASE):
                if not current_section:
                    current_section = {'title': '(без раздела)', 'chapters': []}
                    structure['sections'].append(current_section)

                current_chapter = {
                    'title': text,
                    'index': i,
                    'articles': []
                }
                current_section['chapters'].append(current_chapter)
                continue

            # Статья 1, 2, 3
            if re.match(r'^Статья\s+[\d.]+', text, re.IGNORECASE):
                article = {'title': text, 'index': i}

                if current_chapter:
                    current_chapter['articles'].append(article)
                elif current_section:
                    if 'articles' not in current_section:
                        current_section['articles'] = []
                    current_section['articles'].append(article)

        return structure

    def _extract_decree_structure(self) -> Dict[str, Any]:
        """Структура Постановления: Разделы (римские) + Параграфы + Приложения"""
        structure = {
            'type': 'GOVERNMENT_DECREE',
            'sections': [],
            'appendices': [],
            'total_paragraphs': len(self.paragraphs),
            'total_tables': len(self.doc.tables)
        }

        current_section = None
        in_main_part = False
        section_paras = []

        for i, p in enumerate(self.paragraphs):
            text = p.text.strip()

            # Начало "ПРАВИЛА"
            if 'ПРАВИЛА' in text:
                in_main_part = True
                continue

            # Приложение
            if 'Приложение' in text or 'ПРИЛОЖЕНИЕ' in text:
                if current_section and section_paras:
                    current_section['paragraphs'] = section_paras
                    section_paras = []

                structure['appendices'].append({
                    'title': text,
                    'index': i
                })
                current_section = None
                in_main_part = False
                continue

            # Раздел I. Общие положения
            if in_main_part and re.match(r'^[IVXLCDM]+\.\s+[А-ЯЁ]', text):
                if current_section and section_paras:
                    current_section['paragraphs'] = section_paras
                    section_paras = []

                current_section = {
                    'title': text,
                    'index': i,
                    'paragraphs': []
                }
                structure['sections'].append(current_section)
                continue

            # Параграфы 1., 2., 3.
            if current_section and re.match(r'^\d+\.', text):
                section_paras.append({
                    'number': text.split('.')[0],
                    'text': text[:150],  # Первые 150 символов
                    'index': i
                })

        # Последний раздел
        if current_section:
            current_section['paragraphs'] = section_paras

        return structure

    def _extract_order_structure(self) -> Dict[str, Any]:
        """Структура Приказа: Разделы + Параграфы + Приложения"""
        # Приказы похожи на Постановления
        return self._extract_decree_structure()

    def _extract_simple_structure(self) -> Dict[str, Any]:
        """Простая структура для писем"""
        return {
            'type': self.doc_type,
            'total_paragraphs': len(self.paragraphs),
            'total_tables': len(self.doc.tables),
            'note': 'Структура не определена'
        }

    def save_structure(self, output_path: Path) -> Dict[str, Any]:
        """Сохраняет структуру в файл"""

        structure = self.extract_structure()

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"--- Структура для документа: {self.docx_path.name} ---\n\n")

            # Тип документа
            type_names = {
                'CODE': 'Кодекс / Федеральный закон',
                'GOVERNMENT_DECREE': 'Постановление Правительства РФ',
                'MINISTRY_ORDER': 'Приказ министерства',
                'LETTER': 'Письмо'
            }
            f.write(f"ТИП ДОКУМЕНТА: {type_names.get(structure['type'], structure['type'])}\n\n")

            # Статистика
            f.write(f"СТАТИСТИКА:\n")
            f.write(f"Всего параграфов: {structure.get('total_paragraphs', 0)}\n")
            f.write(f"Всего таблиц: {structure.get('total_tables', 0)}\n")

            if structure['type'] == 'CODE':
                f.write(f"Разделов: {len(structure['sections'])}\n")
                total_chapters = sum(len(s.get('chapters', [])) for s in structure['sections'])
                total_articles = sum(
                    sum(len(ch.get('articles', [])) for ch in s.get('chapters', []))
                    for s in structure['sections']
                )
                f.write(f"Глав: {total_chapters}\n")
                f.write(f"Статей: {total_articles}\n")
            elif structure['type'] in ['GOVERNMENT_DECREE', 'MINISTRY_ORDER']:
                f.write(f"Разделов: {len(structure.get('sections', []))}\n")
                f.write(f"Приложений: {len(structure.get('appendices', []))}\n")
                total_paras = sum(len(s.get('paragraphs', [])) for s in structure.get('sections', []))
                f.write(f"Параграфов в разделах: {total_paras}\n")

            f.write("\n")

            # Основная часть
            if structure['type'] == 'CODE':
                f.write(f"ОСНОВНАЯ ЧАСТЬ:\n\n")
                for section in structure['sections']:
                    f.write(f"{section['title']}\n")

                    # Главы
                    for chapter in section.get('chapters', []):
                        f.write(f"  {chapter['title']}\n")
                        # Статьи
                        for article in chapter.get('articles', []):
                            f.write(f"    {article['title']}\n")

                    # Статьи без глав
                    for article in section.get('articles', []):
                        f.write(f"  {article['title']}\n")

                    f.write("\n")

            elif structure['type'] in ['GOVERNMENT_DECREE', 'MINISTRY_ORDER']:
                if structure.get('sections'):
                    f.write(f"ОСНОВНАЯ ЧАСТЬ (параграфы):\n\n")

                    for section in structure['sections']:
                        f.write(f"{section['title']}\n")
                        for para in section.get('paragraphs', [])[:10]:  # Первые 10
                            f.write(f"  {para['number']}. {para['text']}\n")

                        if len(section.get('paragraphs', [])) > 10:
                            remaining = len(section['paragraphs']) - 10
                            f.write(f"  ... и еще {remaining} параграфов\n")

                        f.write("\n")

                # Приложения
                if structure.get('appendices'):
                    f.write(f"ПРИЛОЖЕНИЯ:\n\n")
                    for app in structure['appendices']:
                        f.write(f"{app['title']} (строка {app['index']})\n")
                    f.write("\n")

            else:
                f.write(f"ОСНОВНАЯ ЧАСТЬ:\n")
                f.write(f"  (структура в формате Раздел/Глава/Статья не найдена)\n\n")
                f.write(f"КОНЕЦ ДОКУМЕНТА:\n")

        logger.info(f"Структура сохранена: {output_path}")
        return structure


def main():
    import sys

    if len(sys.argv) < 2:
        print("Использование: python create_structure.py <имя_файла.docx>")
        print("Пример: python create_structure.py 'ПП РФ № 306.docx'")
        sys.exit(1)

    docx_name = sys.argv[1]
    docs_dir = Path("/home/olga/normativ_docs/Волков/fulldocx")
    docx_path = docs_dir / docx_name

    if not docx_path.exists():
        print(f"Файл не найден: {docx_path}")
        sys.exit(1)

    extractor = UniversalStructureExtractor(docx_path)
    structure_path = docs_dir / f"{docx_path.stem}_structure.txt"

    result = extractor.save_structure(structure_path)

    # Вывод результатов
    print("\n" + "="*80)
    print("РЕЗУЛЬТАТЫ АНАЛИЗА СТРУКТУРЫ")
    print("="*80)
    print(f"Документ: {docx_name}")
    print(f"Тип: {result['type']}")
    print(f"Параграфов: {result.get('total_paragraphs', 0)}")
    print(f"Таблиц: {result.get('total_tables', 0)}")

    if result['type'] == 'CODE':
        print(f"Разделов: {len(result['sections'])}")
    elif result['type'] in ['GOVERNMENT_DECREE', 'MINISTRY_ORDER']:
        print(f"Разделов: {len(result.get('sections', []))}")
        print(f"Приложений: {len(result.get('appendices', []))}")

        print("\nРАЗДЕЛЫ:")
        for section in result.get('sections', []):
            para_count = len(section.get('paragraphs', []))
            print(f"  - {section['title'][:50]}... ({para_count} параграфов)")

    print(f"\nФайл структуры: {structure_path}")


if __name__ == "__main__":
    main()
