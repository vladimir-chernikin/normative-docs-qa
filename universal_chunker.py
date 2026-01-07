#!/usr/bin/env python3
"""
Универсальный чанкер для разных типов нормативных документов
Автоматически определяет тип документа и использует соответствующую стратегию

Поддерживаемые типы документов:
- CODE: Кодексы и Федеральные законы (использует SmartDocumentChunker)
- GOVERNMENT_DECREE: Постановления Правительства (с подпунктами типа 80(1), 80(2))
- MINISTRY_ORDER: Приказы министерств
- LETTER: Письма (простой чанкинг)

Особенности чанкинга постановлений:
- Разбивает по приложениям
- Разбивает по разделам (римские цифры)
- Разбивает по пунктам, включая подпункты (80, 80(1), 80(2) - в одном чанке)
- Ограничение на размер чанка: ~1500 символов
"""

import re
import logging
from pathlib import Path
from typing import List, Dict, Any
from docx import Document
import json

from config import FULLDOCS_DIR, REPORTS_DIR

logger = logging.getLogger(__name__)


class UniversalDocumentChunker:
    """Универсальный чанкер для всех типов документов"""

    def __init__(self, docx_file: Path, structure_file: Path):
        self.docx_file = docx_file
        self.structure_file = structure_file
        self.doc_type = self._detect_document_type()

    def _detect_document_type(self) -> str:
        """Определяет тип документа по названию файла"""

        name = self.docx_file.stem.lower()

        # Кодексы и Федеральные законы (ОБА имеют структуру Глава + Статья)
        if 'кодекс' in name or 'жилищный кодекс' in name:
            return 'CODE'

        if 'фз' in name or '№' in name and 'федеральный закон' in name:
            return 'CODE'  # ФЗ имеют ту же структуру что кодексы!

        # Постановления Правительства
        if 'пп рф' in name or 'постановление' in name:
            return 'GOVERNMENT_DECREE'

        # Приказы
        if 'приказ' in name:
            return 'MINISTRY_ORDER'

        # Письма
        if 'письмо' in name:
            return 'LETTER'

        # По умолчанию - кодекс
        return 'CODE'

    def extract_chunks(self) -> List[Dict[str, Any]]:
        """Извлекает чанки в зависимости от типа документа"""

        logger.info(f"Тип документа: {self.doc_type}")

        if self.doc_type == 'CODE':
            return self._extract_code_chunks()
        elif self.doc_type == 'GOVERNMENT_DECREE':
            return self._extract_government_decree_chunks()
        elif self.doc_type == 'MINISTRY_ORDER':
            return self._extract_ministry_order_chunks()
        else:
            # Для писем и других используем простой чанкинг
            return self._extract_simple_chunks()

    def _extract_code_chunks(self) -> List[Dict[str, Any]]:
        """Чанкер для кодексов (ГК РФ, ЖК РФ)"""
        from smart_chunker import SmartDocumentChunker
        chunker = SmartDocumentChunker(self.docx_file, self.structure_file)
        return chunker.extract_text_with_structure()

    def _extract_government_decree_chunks(self) -> List[Dict[str, Any]]:
        """Чанкер для постановлений Правительства с ограничением размера"""

        MAX_CHUNK_SIZE = 1000  # Максимальный размер чанка (уменьшен для лучшего разбиения)

        # Читаем структуру
        structure_content = self.structure_file.read_text(encoding='utf-8')

        # Читаем DOCX
        doc = Document(str(self.docx_file))
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # Извлекаем приложения из структуры
        apps = self._parse_appendices(structure_content)

        # Разбиваем текст по приложениям и пунктам
        raw_chunks = []
        current_app = None
        current_section = None
        current_chunk_lines = []
        current_main_point = None  # Для отслеживания основного пункта

        for line in full_text:
            # Проверяем приложение
            app_match = re.match(r'Приложение\s+№?\s*\d+', line, re.IGNORECASE)
            if app_match:
                # Сохраняем предыдущий чанк
                if current_chunk_lines:
                    raw_chunks.append({
                        'text': '\n'.join(current_chunk_lines),
                        'metadata': {
                            'document': self.docx_file.stem,
                            'type': 'Постановление',
                            'app': current_app,
                            'section': current_section
                        }
                    })

                current_app = line[:100]
                current_section = None
                current_chunk_lines = [line]
                current_main_point = None  # Сбрасываем пункт
                continue

            # Проверяем раздел (римские цифры)
            section_match = re.match(r'^([IVXLCDM]+)\.\s+(.+)', line)
            if section_match:
                current_section = line[:100]
                if len(current_chunk_lines) > 1:
                    raw_chunks.append({
                        'text': '\n'.join(current_chunk_lines),
                        'metadata': {
                            'document': self.docx_file.stem,
                            'type': 'Постановление',
                            'app': current_app,
                            'section': current_section
                        }
                    })
                current_chunk_lines = [line]
                current_main_point = None  # Сбрасываем пункт
                continue

            # Проверяем пункт (включая подпункты типа 80(1))
            point_match = re.match(r'^(\d+)(\(\d+\))?\.\s+(.+)', line)
            if point_match:
                main_point_number = point_match.group(1)  # Основной номер (80 из 80(1))

                # Если это новый основной пункт - разбиваем чанк
                if main_point_number != current_main_point:
                    # Сохраняем предыдущий чанк
                    if current_chunk_lines:
                        raw_chunks.append({
                            'text': '\n'.join(current_chunk_lines),
                            'metadata': {
                                'document': self.docx_file.stem,
                                'type': 'Постановление',
                                'app': current_app,
                                'section': current_section
                            }
                        })
                    current_main_point = main_point_number
                    current_chunk_lines = [line]
                else:
                    # Это подпункт того же основного пункта - добавляем к текущему чанку
                    current_chunk_lines.append(line)
            else:
                current_chunk_lines.append(line)

        # Добавляем последний чанк
        if current_chunk_lines:
            raw_chunks.append({
                'text': '\n'.join(current_chunk_lines),
                'metadata': {
                    'document': self.docx_file.stem,
                    'type': 'Постановление',
                    'app': current_app,
                    'section': current_section
                }
            })

        # Разбиваем крупные чанки по подпунктам (сохраняя смысл)
        chunks = []
        for chunk in raw_chunks:
            text = chunk['text']
            size = len(text)

            if size <= MAX_CHUNK_SIZE:
                # Нормальный чанк - добавляем как есть
                chunks.append(chunk)
            else:
                # Слишком большой - разбиваем по подпунктам
                logger.info(f"Разбиваем большой чанк ({size} символов) на части")
                sub_chunks = self._split_by_subpoints(text, MAX_CHUNK_SIZE)

                # Рекурсивно проверяем под-чанки - если все еще огромные, разбиваем по предложениям
                for sub_text in sub_chunks:
                    if len(sub_text) > 2000:
                        logger.info(f"Рекурсивно разбиваем под-чанк ({len(sub_text)} символов) по предложениям")
                    final_subs = self._split_by_sentences(sub_text, 1500)
                    for final_text in final_subs:
                        chunks.append({
                            'text': final_text,
                            'metadata': chunk['metadata'].copy()
                        })

        return chunks

    def _split_by_subpoints(self, text: str, max_size: int) -> List[str]:
        """Разбивает крупный чанк по границам подпунктов или параграфов"""

        lines = text.split('\n')
        chunks = []
        current_chunk = []
        current_size = 0

        for line in lines:
            line_size = len(line)

            # Проверяем - это новая логическая единица?
            is_new_unit = (
                re.match(r'^\d+\(\d+\)\.', line) or  # Подпункт 80(1)
                re.match(r'^\d+\.', line) or           # Пункт 80
                re.match(r'^[а-я]\)', line) or         # Буквенный подпункт а)
                (line and line[0].isupper() and len(line) < 200)  # Короткая заглавная строка
            )

            # Если добавление превысит лимит И это новая единица - разбиваем
            if is_new_unit and current_chunk and current_size + line_size > max_size:
                chunks.append('\n'.join(current_chunk))
                current_chunk = [line]
                current_size = line_size
            else:
                current_chunk.append(line)
                current_size += line_size

        # Последний чанк
        if current_chunk:
            chunks.append('\n'.join(current_chunk))

        return chunks

    def _split_by_sentences(self, text: str, max_size: int) -> List[str]:
        """Разбивает чанк по предложениям (для рекурсивного разбиения огромных чанков)"""

        # Разбиваем по предложениям (ищем точку, восклицательный или вопросительный знак)
        sentences = []
        current_sentence = []

        # Разбиваем текст на строки и обрабатываем
        lines = text.split('\n')
        for line in lines:
            # Разбиваем строку на предложения по точкам, восклицательным и вопросительным знакам
            remaining = line
            while remaining:
                # Ищем конец предложения (. ! ?) с последующим пробелом или концом строки
                match = re.search(r'[^.!?]*[.!?](?:\s+|$)', remaining)
                if match:
                    sentence_part = match.group(0).strip()
                    if sentence_part:
                        current_sentence.append(sentence_part)
                    remaining = remaining[match.end():]

                    # Если есть последующая часть и она начинается с маленькой буквы - это продолжение
                    if remaining and remaining[0].islower() and current_sentence:
                        # Продолжаем накапливать
                        continue
                    elif current_sentence:
                        # Конец предложения
                        sentences.append(' '.join(current_sentence))
                        current_sentence = []
                else:
                    # Не найден конец предложения - берем остаток как есть
                    if remaining.strip():
                        current_sentence.append(remaining.strip())
                    break

        # Добавляем последнее предложение
        if current_sentence:
            sentences.append(' '.join(current_sentence))

        # Собираем предложения в чанки
        chunks = []
        current_chunk = []
        current_size = 0

        for sentence in sentences:
            sentence_size = len(sentence)

            # Если добавление превысит лимит и есть уже что-то в чанке - разбиваем
            if current_chunk and current_size + sentence_size > max_size:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = sentence_size
            else:
                current_chunk.append(sentence)
                current_size += sentence_size

        # Последний чанк
        if current_chunk:
            chunks.append(' '.join(current_chunk))

        return chunks if chunks else [text]

    def _extract_ministry_order_chunks(self) -> List[Dict[str, Any]]:
        """Чанкер для приказов министерств"""

        # Читаем DOCX
        doc = Document(str(self.docx_file))
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # Разбиваем по пунктам
        chunks = []
        current_chunk_lines = []
        current_point = None

        for line in full_text:
            # Проверяем пункт
            point_match = re.match(r'^(\d+\.?\d*)\.?\s+(.+)', line)

            if point_match and len(current_chunk_lines) > 3:
                # Сохраняем предыдущий
                chunks.append({
                    'text': '\n'.join(current_chunk_lines),
                    'metadata': {
                        'document': self.docx_file.stem,
                        'type': 'Приказ',
                        'point': current_point
                    }
                })
                current_point = point_match.group(1)
                current_chunk_lines = [line]
            else:
                if not current_point:
                    # Первая строка
                    point_match = re.match(r'^(\d+)\.', line)
                    if point_match:
                        current_point = point_match.group(1)
                current_chunk_lines.append(line)

        # Последний чанк
        if current_chunk_lines:
            chunks.append({
                'text': '\n'.join(current_chunk_lines),
                'metadata': {
                    'document': self.docx_file.stem,
                    'type': 'Приказ',
                    'point': current_point
                }
            })

        return chunks

    def _extract_simple_chunks(self) -> List[Dict[str, Any]]:
        """Простой чанкер для писем и других документов"""

        # Читаем DOCX
        doc = Document(str(self.docx_file))
        full_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # Разбиваем на чанки по 2000 символов
        chunks = []
        chunk_size = 2000
        overlap = 200

        for i in range(0, len(full_text), chunk_size // 10):
            chunk_text = []
            current_length = 0

            for line in full_text[i:]:
                if current_length + len(line) > chunk_size and chunk_text:
                    chunks.append({
                        'text': '\n'.join(chunk_text),
                        'metadata': {
                            'document': self.docx_file.stem,
                            'type': 'Документ'
                        }
                    })
                    chunk_text = []
                    current_length = 0

                chunk_text.append(line)
                current_length += len(line)

            if not chunk_text:
                break

        return chunks

    def _parse_appendices(self, structure_text: str) -> List[str]:
        """Парсит приложения из структуры"""
        apps = []
        for match in re.finditer(r'Приложение\s+№?\s*\d+[:\.]', structure_text, re.IGNORECASE):
            apps.append(match.group(0))
        return apps


# Тестовый запуск
if __name__ == '__main__':
    import sys

    if len(sys.argv) > 1:
        # Передали путь к документу
        docx_path = Path(sys.argv[1])
        structure_path = Path(str(docx_path).replace('.docx', '_structure.txt'))
    else:
        # Тест по умолчанию
        docx_path = FULLDOCS_DIR / "ПП РФ от 03.04.2013 № 290 О минимальном перечне услуг и работ по содержанию общего имущества в МКД.docx"
        structure_path = FULLDOCS_DIR / "ПП РФ от 03.04.2013 № 290 О минимальном перечне услуг и работ по содержанию общего имущества в МКД_structure.txt"

    logger.info(f"Обработка: {docx_path.name}")
    logger.info(f"Структура: {structure_path.name}")

    chunker = UniversalDocumentChunker(docx_path, structure_path)
    chunks = chunker.extract_chunks()

    logger.info(f"Создано {len(chunks)} чанков")

    # Показываем примеры
    logger.info("ПРИМЕРЫ ЧАНКОВ:")
    logger.info("=" * 80)
    for i, chunk in enumerate(chunks[:3], 1):
        logger.info(f"\nЧАНК #{i}")
        logger.info(f"Документ: {chunk['metadata']['document']}")
        logger.info(f"Тип: {chunk['metadata'].get('type', 'Неизвестно')}")
        logger.info(f"Метаданные: {chunk['metadata']}")
        logger.info(f"Текст (первые 200 символов): {chunk['text'][:200]}...")
        logger.info("-" * 80)

    # Сохраняем
    output_file = REPORTS_DIR / f"test_chunks_{chunker.doc_type.lower()}.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(chunks, f, ensure_ascii=False, indent=2)

    logger.info(f"Сохранено в: {output_file}")
